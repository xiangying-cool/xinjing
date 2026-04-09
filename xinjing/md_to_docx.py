"""Convert 第4章_系统实现.md to Word (.docx)"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

MD_PATH = Path(__file__).parent / "第4章_系统实现.md"
OUT_PATH = Path(__file__).parent / "第4章_系统实现.docx"

text = MD_PATH.read_text(encoding="utf-8")
lines = text.splitlines()

doc = Document()

# ── 页面设置 ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.18)
section.right_margin  = Cm(3.18)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ── 默认正文样式 ──────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = '宋体'
style_normal.font.size = Pt(12)
style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_heading_style(para, level):
    """Apply heading style with font settings."""
    sizes = {1: 18, 2: 16, 3: 14, 4: 13}
    bold_levels = {1, 2, 3, 4}
    run = para.runs[0] if para.runs else para.add_run()
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = level in bold_levels
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_heading(text_content, level):
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
    para = doc.add_heading(text_content, level=level)
    style_name = style_map.get(level, 'Heading 1')
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass
    # Override font
    for run in para.runs:
        run.font.size = Pt({1:18,2:16,3:14,4:13}.get(level,12))
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return para

def add_paragraph(text_content, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = Pt(20)
    _apply_inline(para, text_content)
    if indent:
        para.paragraph_format.first_line_indent = Pt(24)
    return para

def _apply_inline(para, text_content):
    """Handle **bold**, `code`, and plain text inline."""
    # Split on **bold** and `code`
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\$[^$]+\$)')
    parts = pattern.split(text_content)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif part.startswith('$') and part.endswith('$'):
            # Math: just render as-is in italic
            run = para.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Cambria Math'
        else:
            if part:
                run = para.add_run(part)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_table(header_row, data_rows):
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Table Grid'
    # Header
    hrow = table.rows[0]
    for i, cell_text in enumerate(header_row):
        cell = hrow.cells[i]
        cell.text = cell_text
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(cell_text)
        run.bold = True
        run.font.name = '黑体'
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header cell shading
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E2F3')
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(data_rows):
        trow = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = trow.cells[ci]
            cell.text = ''
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            _apply_inline(para, cell_text)
            for run in para.runs:
                run.font.size = Pt(10.5)
    doc.add_paragraph()

def add_code_block(code_lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    # Light gray shading via paragraph border/shading would need XML; use a simple style
    run = para.add_run('\n'.join(code_lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # Add shading to paragraph
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

def add_bullet(text_content, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    para.paragraph_format.space_after = Pt(4)
    _apply_inline(para, text_content)
    for run in para.runs:
        if not run.font.name or run.font.name == 'Calibri':
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)

# ── 解析状态机 ────────────────────────────────────────────────────────────────
i = 0
in_code = False
code_lines = []
in_table = False
table_header = []
table_data = []

while i < len(lines):
    line = lines[i]

    # Code block
    if line.strip().startswith('```'):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            in_code = False
            add_code_block(code_lines)
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # Table
    if line.strip().startswith('|'):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if not in_table:
            in_table = True
            table_header = cells
            table_data = []
        else:
            # Skip separator row (---|---|---)
            if all(re.match(r'^[-: ]+$', c) for c in cells if c):
                i += 1
                continue
            table_data.append(cells)
        i += 1
        # Peek next line
        if i >= len(lines) or not lines[i].strip().startswith('|'):
            add_table(table_header, table_data)
            in_table = False
        continue

    in_table = False

    # Headings
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        add_heading(m.group(2), level)
        i += 1
        continue

    # Horizontal rule
    if re.match(r'^---+$', line.strip()):
        doc.add_paragraph('─' * 40)
        i += 1
        continue

    # Blockquote
    if line.startswith('>'):
        content = line.lstrip('> ').strip()
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1)
        para.paragraph_format.space_after = Pt(4)
        _apply_inline(para, content)
        for run in para.runs:
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.size = Pt(11)
        i += 1
        continue

    # Bullet list
    m_bullet = re.match(r'^(\s*)-\s+(.*)', line)
    if m_bullet:
        indent_level = len(m_bullet.group(1)) // 2
        add_bullet(m_bullet.group(2), indent_level)
        i += 1
        continue

    # Numbered list
    m_num = re.match(r'^\s*\d+\.\s+(.*)', line)
    if m_num:
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(4)
        _apply_inline(para, m_num.group(1))
        for run in para.runs:
            if not run.font.name or run.font.name == 'Calibri':
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
        i += 1
        continue

    # Empty line
    if not line.strip():
        i += 1
        continue

    # Normal paragraph
    add_paragraph(line.strip(), indent=True)
    i += 1

doc.save(str(OUT_PATH))
print(f"Saved: {OUT_PATH}")
